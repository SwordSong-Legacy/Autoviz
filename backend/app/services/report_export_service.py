"""Report export service: PDF, DOCX, MD with embedded visualization images.

Produces professional reports: image centered on top, insight text below.
"""

from __future__ import annotations

import base64
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Image as RlImage,
)
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

# Max image width for report layout
IMG_WIDTH_CM = 14
IMG_WIDTH_MM = 140


def _normalize_detail(d: Any) -> tuple[str, list[int]]:
    """Normalize detail to (text, viz_refs)."""
    if isinstance(d, dict):
        text = d.get("text", "")
        viz_refs = d.get("viz_refs") or []
        if isinstance(viz_refs, list):
            viz_refs = [int(x) for x in viz_refs if isinstance(x, (int, float))]
        else:
            viz_refs = []
        return text, viz_refs
    return "", []


def _iter_finding_blocks(
    key_findings: list[Any],
) -> list[tuple[str, str, list[tuple[str, list[int]]]]]:
    """Yield (theme, summary, [(detail_text, viz_refs), ...]) for each theme."""
    blocks: list[tuple[str, str, list[tuple[str, list[int]]]]] = []
    for kf in key_findings or []:
        if isinstance(kf, dict):
            theme = kf.get("theme", "")
            summary = kf.get("summary", "")
            details_raw = kf.get("details") or []
        else:
            continue
        details = [_normalize_detail(d) for d in details_raw]
        blocks.append((theme, summary, details))
    return blocks


def export_markdown(
    dataset_overview: str,
    key_findings: list[Any],
    relationship_analysis: list[Any],
    suggestions: list[str],
    viz_images: list[bytes],
    filename_base: str = "report",
) -> tuple[bytes, str]:
    """Export report as Markdown with base64-embedded images.

    Returns (content_bytes, suggested_filename).
    """
    lines: list[str] = []
    lines.append("# Dataset Overview\n")
    lines.append(dataset_overview)
    lines.append("\n# Key Findings\n")

    shown_viz: set[int] = set()
    for theme, summary, details in _iter_finding_blocks(key_findings):
        lines.append(f"## {theme}\n")
        lines.append(summary)
        if details:
            lines.append("")
            for text, viz_refs in details:
                if viz_refs:
                    for idx in viz_refs:
                        if (
                            idx not in shown_viz
                            and 1 <= idx <= len(viz_images)
                            and viz_images[idx - 1]
                        ):
                            shown_viz.add(idx)
                            b64 = base64.b64encode(viz_images[idx - 1]).decode("ascii")
                            lines.append(f"![Visualization {idx}](data:image/png;base64,{b64})\n")
                lines.append(f"- {text}")
                lines.append("")
        lines.append("")

    lines.append("# Relationship Analysis\n")
    for ra in relationship_analysis or []:
        if isinstance(ra, dict):
            vars_ = ra.get("variables") or []
            desc = ra.get("description", "")
            lines.append(f"- **{', '.join(vars_)}**: {desc}")
    lines.append("\n# Suggestions\n")
    for s in suggestions or []:
        lines.append(f"- {s}")

    content = "\n".join(lines)
    return content.encode("utf-8"), f"{filename_base}.md"


def export_docx(
    dataset_overview: str,
    key_findings: list[Any],
    relationship_analysis: list[Any],
    suggestions: list[str],
    viz_images: list[bytes],
    filename_base: str = "report",
) -> tuple[bytes, str]:
    """Export report as DOCX with embedded images. Image above, insight below."""
    doc = Document()
    doc.add_heading("Dataset Overview", level=0)
    doc.add_paragraph(dataset_overview)

    doc.add_heading("Key Findings", level=0)
    shown_viz: set[int] = set()
    for theme, summary, details in _iter_finding_blocks(key_findings):
        doc.add_heading(theme, level=1)
        doc.add_paragraph(summary)
        for text, viz_refs in details:
            if viz_refs:
                for idx in viz_refs:
                    if idx not in shown_viz and 1 <= idx <= len(viz_images) and viz_images[idx - 1]:
                        shown_viz.add(idx)
                        img_stream = BytesIO(viz_images[idx - 1])
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(img_stream, width=Cm(IMG_WIDTH_CM))
                doc.add_paragraph()
            doc.add_paragraph(text, style="List Bullet")
        doc.add_paragraph()

    doc.add_heading("Relationship Analysis", level=0)
    for ra in relationship_analysis or []:
        if isinstance(ra, dict):
            vars_ = ra.get("variables") or []
            desc = ra.get("description", "")
            doc.add_paragraph(f"{', '.join(vars_)}: {desc}", style="List Bullet")

    doc.add_heading("Suggestions", level=0)
    for s in suggestions or []:
        doc.add_paragraph(s, style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue(), f"{filename_base}.docx"


def export_pdf(
    dataset_overview: str,
    key_findings: list[Any],
    relationship_analysis: list[Any],
    suggestions: list[str],
    viz_images: list[bytes],
    filename_base: str = "report",
) -> tuple[bytes, str]:
    """Export report as PDF with embedded images. Image above, insight below."""
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )
    styles = getSampleStyleSheet()
    story: list = []

    def add_heading(text: str, level: int = 0) -> None:
        size = 18 - level * 2
        style = ParagraphStyle(
            name=f"Heading{level}",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=size,
            spaceAfter=6,
        )
        story.append(Paragraph(text, style))
        story.append(Spacer(1, 4 * mm))

    def add_para(text: str, bullet: bool = False) -> None:
        escaped = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if bullet:
            escaped = "• " + escaped
        story.append(Paragraph(escaped, styles["Normal"]))
        story.append(Spacer(1, 2 * mm))

    add_heading("Dataset Overview")
    add_para(dataset_overview)

    add_heading("Key Findings")
    shown_viz: set[int] = set()
    for theme, summary, details in _iter_finding_blocks(key_findings):
        add_heading(theme, level=1)
        add_para(summary)
        for text, viz_refs in details:
            if viz_refs:
                for idx in viz_refs:
                    if idx not in shown_viz and 1 <= idx <= len(viz_images) and viz_images[idx - 1]:
                        shown_viz.add(idx)
                        img = RlImage(
                            BytesIO(viz_images[idx - 1]),
                            width=IMG_WIDTH_MM,
                            height=IMG_WIDTH_MM * 0.6,
                        )
                        img.hAlign = "CENTER"
                        story.append(img)
                        story.append(Spacer(1, 4 * mm))
            add_para(text, bullet=True)
        story.append(Spacer(1, 4 * mm))

    add_heading("Relationship Analysis")
    for ra in relationship_analysis or []:
        if isinstance(ra, dict):
            vars_ = ra.get("variables") or []
            desc = ra.get("description", "")
            add_para(f"{', '.join(vars_)}: {desc}", bullet=True)

    add_heading("Suggestions")
    for s in suggestions or []:
        add_para(s, bullet=True)

    doc.build(story)
    buf.seek(0)
    return buf.getvalue(), f"{filename_base}.pdf"
